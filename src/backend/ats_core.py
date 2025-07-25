"""
Enhanced ATS Core Module for CV-to-JD Matching
Provides comprehensive parsing, NLP processing, and matching algorithms
Framework-agnostic design for maximum flexibility
"""

import re
import io
from typing import List, Dict, Optional, Tuple
from docx import Document
from pypdf import PdfReader
import spacy
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize, sent_tokenize
from spacy.matcher import PhraseMatcher
import logging
import spacy.cli # Import spacy.cli for programmatic downloads

# Configure logging for better visibility of issues
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Global spaCy model - loaded once for efficiency
nlp = None

def initialize_nlp():
    """Initialize spaCy model and NLTK data"""
    global nlp
    try:
        # Try to load 'en_core_web_sm' first, as requested in previous instructions
        nlp = spacy.load("en_core_web_sm")
        logging.info("Successfully loaded spaCy model 'en_core_web_sm'.")
    except OSError:
        # If model not found, attempt to download it
        logging.warning("spaCy model 'en_core_web_sm' not found. Attempting to download it.")
        try:
            spacy.cli.download("en_core_web_sm")
            nlp = spacy.load("en_core_web_sm")
            logging.info("Successfully downloaded and loaded spaCy model 'en_core_web_sm'.")
        except Exception as e:
            logging.error(f"Failed to download or load spaCy model 'en_core_web_sm' even after attempt: {e}")
            nlp = None # Ensure nlp is None if loading failed
            return False # Indicate failure
    except Exception as e:
        logging.error(f"Failed to initialize NLP models due to an unexpected error: {e}")
        nlp = None # Ensure nlp is None if loading failed
        return False # Indicate failure

    # Download required NLTK data
    try:
        nltk.download('stopwords', quiet=True)
        nltk.download('punkt', quiet=True)
        logging.info("Successfully downloaded NLTK data (stopwords, punkt).")
        return True # Indicate overall success
    except Exception as e:
        logging.error(f"Failed to download NLTK data: {e}")
        return False # Indicate failure


# Common professional skills dictionary for enhanced matching
PROFESSIONAL_SKILLS = {
    # Programming Languages
    'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'php', 'ruby', 'go', 'rust',
    'swift', 'kotlin', 'scala', 'r', 'matlab', 'sql', 'html', 'css', 'bash', 'powershell',
    
    # Frameworks & Libraries
    'react', 'angular', 'vue', 'node.js', 'express', 'django', 'flask', 'spring', 'laravel',
    'rails', 'asp.net', 'jquery', 'bootstrap', 'tailwind', 'tensorflow', 'pytorch', 'pandas',
    'numpy', 'scikit-learn', 'keras', 'opencv',
    
    # Technologies & Tools
    'docker', 'kubernetes', 'jenkins', 'git', 'github', 'gitlab', 'aws', 'azure', 'gcp',
    'terraform', 'ansible', 'chef', 'puppet', 'vagrant', 'nginx', 'apache', 'redis',
    'elasticsearch', 'mongodb', 'postgresql', 'mysql', 'oracle', 'cassandra', 'kafka', 'hadoop', 'spark',
    
    # Methodologies & Concepts
    'agile', 'scrum', 'kanban', 'devops', 'ci/cd', 'tdd', 'bdd', 'microservices',
    'rest api', 'graphql', 'soap', 'oauth', 'jwt', 'ssl', 'https', 'encryption',
    'machine learning', 'deep learning', 'artificial intelligence', 'data science',
    'big data', 'cloud computing', 'blockchain', 'iot', 'cybersecurity', 'data analysis', 'data visualization',
    
    # Soft Skills
    'leadership', 'project management', 'team management', 'communication', 'problem solving',
    'analytical thinking', 'strategic planning', 'stakeholder management', 'mentoring',
    'cross-functional collaboration', 'client management', 'vendor management', 'adaptability', 'creativity',
    'critical thinking', 'negotiation', 'time management', 'attention to detail', 'innovation'
}

# Job seniority levels mapping
SENIORITY_LEVELS = {
    'intern': 0, 'trainee': 0, 'entry': 1, 'junior': 2, 'associate': 3,
    'mid': 4, 'senior': 5, 'lead': 6, 'principal': 7, 'staff': 7,
    'architect': 8, 'director': 9, 'vp': 10, 'cto': 11, 'ceo': 12
}

def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    """Extract text from PDF bytes with robust error handling"""
    try:
        pdf_file = io.BytesIO(pdf_bytes)
        reader = PdfReader(pdf_file)
        text = ""
        
        for page in reader.pages:
            try:
                text += page.extract_text() + "\n"
            except Exception as e:
                logging.warning(f"Failed to extract text from PDF page: {e}")
                continue
                
        return text.strip()
    except Exception as e:
        logging.error(f"Failed to extract text from PDF: {e}")
        return ""

def extract_text_from_docx(docx_bytes: bytes) -> str:
    """Extract text from DOCX bytes with robust error handling"""
    try:
        docx_file = io.BytesIO(docx_bytes)
        doc = Document(docx_file)
        text = ""
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
            
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
                
        return text.strip()
    except Exception as e:
        logging.error(f"Failed to extract text from DOCX: {e}")
        return ""

def extract_text_from_txt(txt_bytes: bytes) -> str:
    """Extract text from TXT bytes with robust encoding handling"""
    try:
        # Try different encodings
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                return txt_bytes.decode(encoding).strip()
            except UnicodeDecodeError:
                continue
                
        # If all encodings fail, use utf-8 with error handling
        return txt_bytes.decode('utf-8', errors='ignore').strip()
    except Exception as e:
        logging.error(f"Failed to extract text from TXT: {e}")
        return ""

def clean_text(text: str) -> str:
    """Clean and normalize text for processing"""
    try:
        # Convert to lowercase
        text = text.lower()
        
        # Remove extra whitespace and normalize
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep important punctuation
        text = re.sub(r'[^\w\s\-\.\,\(\)#]', ' ', text) # Keep hash for C#
        
        # Remove digits that are standalone (keep version numbers, etc.)
        text = re.sub(r'\b\d+\b', ' ', text)
        
        # Remove stopwords
        try:
            # Ensure stopwords are downloaded
            stop_words = set(stopwords.words('english'))
            word_tokens = word_tokenize(text)
            filtered_text = [word for word in word_tokens if word not in stop_words and len(word) > 2]
            text = ' '.join(filtered_text)
        except Exception as e:
            logging.warning(f"NLTK stopwords/punkt data not available or error during tokenization: {e}. Skipping stopword removal.")
            # Fallback if NLTK data not available or error
            pass
            
        return text.strip()
    except Exception as e:
        logging.error(f"Failed to clean text: {e}")
        return text

def extract_skills(text: str) -> List[str]:
    """Extract skills using enhanced NLP and predefined skill matching"""
    global nlp
    if not nlp:
        # Attempt to initialize if not already successful
        if not initialize_nlp():
            logging.warning("NLP model not loaded, skill extraction will be limited to keyword matching.")
            # Fallback to basic keyword matching if NLP not available
            found_skills = set()
            text_lower = text.lower()
            for skill in PROFESSIONAL_SKILLS:
                if skill in text_lower:
                    found_skills.add(skill)
            return list(found_skills)

    skills = set()
    
    try:
        # Use spaCy for entity recognition and phrase matching
        doc = nlp(text.lower())
        
        # Extract noun phrases and entities that might be skills
        for token in doc:
            if token.pos_ in ['NOUN', 'PROPN'] and len(token.text) > 2:
                if token.text in PROFESSIONAL_SKILLS:
                    skills.add(token.text)
        
        # Extract multi-word skills using phrase matching
        matcher = PhraseMatcher(nlp.vocab, attr="LOWER")
        patterns = [nlp(skill) for skill in PROFESSIONAL_SKILLS if len(skill.split()) > 1]
        if patterns: # Only add if there are patterns to avoid empty matcher issues
            matcher.add("SKILLS", patterns)
        
        matches = matcher(doc)
        for match_id, start, end in matches:
            skill = doc[start:end].text.lower()
            skills.add(skill)
        
        # Pattern-based skill extraction for common formats
        skill_patterns = [
            r'\b(?:proficient|experienced|expert|skilled)\s+(?:in|with)\s+([a-zA-Z0-9\s\-\.#]+)', # Added 0-9 and # for tech skills
            r'\b(?:knowledge|experience)\s+(?:of|in|with)\s+([a-zA-Z0-9\s\-\.#]+)',
            r'\b(?:using|worked with|utilized)\s+([a-zA-Z0-9\s\-\.#]+)',
            r'(?:skills|technologies|proficiencies):\s*([a-zA-Z0-9\s\-\.,#]+)', # for colon-separated lists
        ]
        
        for pattern in skill_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                # Split by common delimiters if it's a list
                potential_skills_str = match.group(1).strip()
                potential_skills_list = re.split(r'[,/;]', potential_skills_str) # Split by comma, slash, or semicolon
                
                for potential_skill in potential_skills_list:
                    potential_skill = potential_skill.strip().lower()
                    # Check if extracted skill or part of it is in PROFESSIONAL_SKILLS
                    if potential_skill in PROFESSIONAL_SKILLS:
                        skills.add(potential_skill)
                    else:
                        # Try to match sub-strings if the full match isn't in PROFESSIONAL_SKILLS
                        for defined_skill in PROFESSIONAL_SKILLS:
                            if defined_skill in potential_skill and len(defined_skill) > 2:
                                skills.add(defined_skill)
        
        return list(skills)
        
    except Exception as e:
        logging.error(f"Failed to extract skills with NLP: {e}")
        # Fallback if NLP processing fails unexpectedly
        found_skills = set()
        text_lower = text.lower()
        for skill in PROFESSIONAL_SKILLS:
            if skill in text_lower:
                found_skills.add(skill)
        return list(found_skills)

def extract_experience(text: str) -> Dict:
    """Extract years of experience and seniority level"""
    try:
        experience_info = {
            "years_of_experience": 0,
            "seniority_level": "Entry-Level"
        }
        
        # Extract years of experience using various patterns
        year_patterns = [
            r'(\d+\.?\d*)\+?\s*(?:years?|yrs?|yr)\s*(?:of\s*)?(?:experience|exp|ex)', # Added float support
            r'(?:experience|exp|ex).*?(\d+\.?\d*)\+?\s*(?:years?|yrs?|yr)',
            r'(\d+\.?\d*)\+?\s*(?:years?|yrs?|yr)\s*(?:in|with)',
            r'over\s+(\d+\.?\d*)\s*(?:years?|yrs?)',
            r'more than\s+(\d+\.?\d*)\s*(?:years?|yrs?)',
            r'(\d+)\s*-\s*(\d+)\s*(?:years?|yrs?)' # Range like "2-5 years"
        ]
        
        years = []
        for pattern in year_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                try:
                    if len(match.groups()) == 2 and match.group(2) is not None: # For range pattern
                        years.append(float(match.group(1)))
                        years.append(float(match.group(2)))
                    else:
                        years.append(float(match.group(1)))
                except ValueError:
                    continue
        
        if years:
            experience_info["years_of_experience"] = max(years)
        
        # Extract seniority from job titles and keywords
        seniority_keywords = {
            'intern': ['intern', 'internship', 'trainee'],
            'junior': ['junior', 'associate', 'entry level', 'entry-level'],
            'mid': ['mid level', 'mid-level', 'intermediate'],
            'senior': ['senior', 'sr.', 'sr '],
            'lead': ['lead', 'team lead', 'tech lead', 'technical lead'],
            'principal': ['principal', 'staff', 'expert'],
            'architect': ['architect', 'solution architect', 'system architect'],
            'director': ['director', 'head of', 'vp', 'vice president'],
            'cto': ['cto', 'chief technology officer', 'chief technical officer']
        }
        
        detected_levels = []
        # Check text for explicit seniority keywords
        for level, keywords in seniority_keywords.items():
            for keyword in keywords:
                if keyword in text.lower():
                    detected_levels.append(level)
        
        # Determine highest seniority level found
        if detected_levels:
            # Get the highest seniority level based on predefined order
            highest_level = max(detected_levels, key=lambda x: SENIORITY_LEVELS.get(x, -1))
            experience_info["seniority_level"] = highest_level.replace('-', ' ').title() # Format nicely
        
        # If no explicit seniority keyword is found, infer from years of experience
        elif experience_info["years_of_experience"] >= 8:
            experience_info["seniority_level"] = "Principal" # More appropriate for 8+
        elif experience_info["years_of_experience"] >= 5:
            experience_info["seniority_level"] = "Senior"
        elif experience_info["years_of_experience"] >= 2:
            experience_info["seniority_level"] = "Mid"
        elif experience_info["years_of_experience"] >= 0: # Catch freshers explicitly
             experience_info["seniority_level"] = "Entry-Level"
        
        return experience_info
        
    except Exception as e:
        logging.error(f"Failed to extract experience: {e}")
        return {"years_of_experience": 0, "seniority_level": "Entry-Level"}

def extract_ctc_info(text: str) -> Dict:
    """Extract CTC/Salary information with robust parsing"""
    try:
        ctc_info = {}
        
        # Currency patterns
        currency_patterns = {
            'USD': [r'\$', r'usd', r'dollars?', r'a\.?k\.?a\.?', r'per annum', r'yearly', r'annually'], # Added AKA and yearly
            'EUR': [r'€', r'eur', r'euros?'],
            'GBP': [r'£', r'gbp', r'pounds?'],
            'INR': [r'₹', r'inr', r'rupees?', r'lpa', r'lakhs?', r'crores?', r'lac'] # Added 'lac'
        }
        
        # Salary extraction patterns (prioritize patterns with explicit currency/unit)
        salary_patterns = [
            r'(\d+(?:[.,]\d+)?)\s*(?:to|-)?\s*(\d+(?:[.,]\d+)?)?\s*(lpa|k|thousand|lakhs?|crores?|usd|eur|gbp|inr|\$|€|£|₹)\s*(?:p\.?a\.?|per annum|annually|yearly)?', # Explicit unit/currency
            r'(?:ctc|salary|compensation|expected|package|pay|remuneration)\s*(?:range)?\s*(?:of|from)?\s*(?:(?:about|up to|min|max)?\s*)?(?:(?:rs\.?|₹|\$|€|£)?\s*(\d+(?:[.,]\d+)?)\s*(?:(?:to|-)\s*(?:rs\.?|₹|\$|€|£)?\s*(\d+(?:[.,]\d+)?))?)?\s*(lpa|k|thousand|lakhs?|crores?|usd|eur|gbp|inr)?', # Keywords and numbers
            r'(?:rs\.?|₹|\$|€|£)\s*(\d+(?:[.,]\d+)?)\s*(?:to|-)?\s*(?:rs\.?|₹|\$|€|£)?\s*(\d+(?:[.,]\d+)?)?', # Currency symbol first
            r'(\d+(?:[.,]\d+)?)\s*(?:lpa|k|thousand|lakhs?|crores?|lac)' # Number followed by unit
        ]
        
        text_lower = text.lower()
        
        for pattern in salary_patterns:
            matches = re.finditer(pattern, text_lower, re.IGNORECASE)
            for match in matches:
                try:
                    # Safely get groups, handling optional ones
                    min_val_str = match.group(1).replace(',', '') if match.group(1) else None
                    max_val_str = match.group(2).replace(',', '') if match.group(2) else min_val_str # If group 2 is None, use group 1
                    
                    # Group 3 is often the unit (lpa, k, usd, etc.)
                    unit = match.group(3).lower() if len(match.groups()) >= 3 and match.group(3) else ''

                    if not min_val_str: # Must have at least a first number
                        continue
                    
                    min_val = float(min_val_str)
                    max_val = float(max_val_str) if max_val_str else min_val
                    
                    # Normalize values based on unit
                    multiplier = 1
                    detected_currency = 'USD' # Default currency if not specified
                    
                    if 'lpa' in unit or 'lakhs' in unit or 'lac' in unit:
                        multiplier = 100000
                        detected_currency = 'INR'
                    elif 'crore' in unit:
                        multiplier = 10000000
                        detected_currency = 'INR'
                    elif 'k' in unit or 'thousand' in unit:
                        multiplier = 1000
                    
                    # Detect currency from symbols or specific keywords in the matched snippet
                    matched_snippet = match.group(0).lower()
                    for curr, patterns in currency_patterns.items():
                        for curr_pattern in patterns:
                            if re.search(curr_pattern, matched_snippet):
                                detected_currency = curr
                                break
                        if detected_currency != 'USD': # If a specific currency is found, break outer loop
                            break

                    ctc_info = {
                        'min_value': min_val * multiplier,
                        'max_value': max_val * multiplier,
                        'currency': detected_currency,
                        'original_text': match.group(0)
                    }
                    return ctc_info # Return as soon as a valid CTC is found
                except (ValueError, AttributeError) as e:
                    logging.debug(f"Skipping CTC match due to parsing error: {e}, text: {match.group(0)}")
                    continue
        
        return ctc_info # Return empty dict if no CTC info found
        
    except Exception as e:
        logging.error(f"Failed to extract CTC info: {e}")
        return {} # Return empty dict on error

def extract_academic_info(text: str) -> Dict:
    """Extract academic achievements and qualifications from CV"""
    try:
        academic_info = {
            "degrees": [],
            "universities": [],
            "publications": [],
            "awards": []
        }
        
        # Degree patterns
        degree_patterns = [
            r'\b(?:bachelor|b\.?tech|b\.?sc|b\.?a|b\.?e|be|btech|bsc|ba)\b',
            r'\b(?:master|m\.?tech|m\.?sc|m\.?a|m\.?e|me|mtech|msc|ma|mba)\b',
            r'\b(?:phd|ph\.?d|doctorate|doctoral)\b',
            r'\b(?:diploma|certificate)\b'
        ]
        
        for pattern in degree_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                # Extract surrounding context for full degree name
                start = max(0, match.start() - 50)
                end = min(len(text), match.end() + 50)
                context = text[start:end].strip()
                academic_info["degrees"].append(context)
        
        # University patterns
        university_patterns = [
            r'\b(?:university|college|institute|school)\s+of\s+[\w\s]+',
            r'[\w\s]+\s+(?:university|college|institute|iit|iim|nit)',
            r'\b(?:iit|iim|nit|iisc|bits)\s+[\w\s]*',
            r'\b(?:[A-Z][a-zA-Z\s,]+\s(?:University|College|Institute|School))\b', # Better capture of names
            r'\b(?:indian\s+institute\s+of\s+technology|indian\s+institute\s+of\s+management)\b',
        ]
        
        for pattern in university_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                academic_info["universities"].append(match.group(0).strip())
        
        # Publication patterns
        publication_patterns = [
            r'(?:published|publication|paper|journal|conference).*?(?:in|at)\s+[\w\s\-]+',
            r'(?:author|co-author|authored).*?(?:paper|article|publication)',
            r'(?:ieee|acm|springer|elsevier|nature|science|cvpr|iccv|eccv|neurips|icml|aaai|acl|emnlp|naacl|ijcai|kdd|sigir).*?(?:conference|journal|symposium|workshop)', # More specific pub terms
            r'doi:\s*[\w\.\/]+' # DOI identifier
        ]
        
        for pattern in publication_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                academic_info["publications"].append(match.group(0).strip())
        
        # Awards patterns
        award_patterns = [
            r'(?:award|prize|recognition|honor|scholarship|fellowship)',
            r'(?:dean\'?s list|honor roll|magna cum laude|summa cum laude|cum laude|valedictorian|salutatorian)',
            r'(?:gold medal|silver medal|bronze medal|first class|distinction|gpa:\s*[\d\.]+)', # GPA added
            r'(?:best\s+paper|best\s+poster|best\s+thesis|hackathon\s+winner|top\s+\d+%)' # Competitive achievements
        ]
        
        for pattern in award_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                # Extract surrounding context
                start = max(0, match.start() - 30)
                end = min(len(text), match.end() + 30)
                context = text[start:end].strip()
                academic_info["awards"].append(context)
        
        # Remove duplicates
        for key in academic_info:
            academic_info[key] = list(set(academic_info[key]))
        
        return academic_info
        
    except Exception as e:
        logging.error(f"Failed to extract academic info: {e}")
        return {"degrees": [], "universities": [], "publications": [], "awards": []}

def extract_key_responsibilities_from_jd(text: str) -> List[str]:
    """Extract key responsibilities and requirements from Job Description"""
    global nlp
    if not nlp:
        if not initialize_nlp():
            logging.warning("NLP model not loaded, responsibility extraction will be limited.")
            # Fallback for responsibility extraction if NLP fails
            sentences = sent_tokenize(text)
            responsibilities = []
            responsibility_indicators = [
                'responsibilities:', 'duties:', 'will be responsible for', 'key responsibilities include',
                'you will', 'we are looking for someone who will', 'able to'
            ]
            for sentence in sentences:
                if any(indicator in sentence.lower() for indicator in responsibility_indicators) or re.match(r'^\s*[\-\*\•\d+\.\)]\s*', sentence.strip()):
                    clean_resp = re.sub(r'^\s*[\-\*\•\d+\.\)]\s*', '', sentence.strip())
                    if len(clean_resp) > 20 and len(clean_resp) < 200:
                        responsibilities.append(clean_resp)
            return list(set(responsibilities))[:10]
    
    try:
        responsibilities = []
        
        # Common responsibility verbs and patterns
        responsibility_verbs = [
            'develop', 'design', 'implement', 'create', 'build', 'maintain',
            'manage', 'lead', 'coordinate', 'oversee', 'supervise', 'direct',
            'analyze', 'optimize', 'improve', 'enhance', 'troubleshoot',
            'collaborate', 'work with', 'partner with', 'communicate',
            'ensure', 'deliver', 'execute', 'perform', 'conduct',
            'architect', 'define', 'integrate', 'test', 'deploy', 'monitor' # Added more tech verbs
        ]
        
        # Split text into sentences
        sentences = sent_tokenize(text)
        
        for sentence in sentences:
            sentence_lower = sentence.lower()
            
            # Check if sentence contains responsibility indicators
            has_responsibility_verb = any(f" {verb}" in f" {sentence_lower}" for verb in responsibility_verbs) # Ensure whole word match
            
            # Look for bullet points or numbered lists
            is_bullet_point = re.match(r'^\s*[\-\*\•\d+\.\)]\s*', sentence.strip())
            
            # Look for requirement indicators
            has_requirement_indicator = any(indicator in sentence_lower for indicator in [
                'required', 'must have', 'should have', 'responsible for',
                'duties include', 'key responsibilities', 'main tasks', 'you will',
                'ability to', 'experience in', 'demonstrated ability'
            ])
            
            if (has_responsibility_verb or is_bullet_point or has_requirement_indicator) and len(sentence.strip()) > 20:
                # Use spaCy to get the main clause/root if available for better extraction
                if nlp:
                    doc_sentence = nlp(sentence.strip())
                    # Try to find the main verb and its direct object for a more concise responsibility
                    main_clause = []
                    # Heuristic: Find root verb and its direct object/complement
                    root_token = None
                    for token in doc_sentence:
                        if token.dep_ == 'ROOT':
                            root_token = token
                            break
                    
                    if root_token:
                        main_clause.append(root_token.text)
                        for child in root_token.children:
                            if child.dep_ in ['dobj', 'attr', 'acomp', 'xcomp', 'prep', 'npadvmod', 'compound', 'amod', 'nmod']:
                                main_clause.append(child.text)
                    
                    if main_clause:
                        clean_resp = ' '.join(main_clause).strip()
                    else:
                        clean_resp = re.sub(r'^\s*[\-\*\•\d+\.\)]\s*', '', sentence.strip())
                else:
                    clean_resp = re.sub(r'^\s*[\-\*\•\d+\.\)]\s*', '', sentence.strip())
                
                clean_resp = clean_resp.strip()
                
                if len(clean_resp) > 15 and len(clean_resp) < 200:
                    responsibilities.append(clean_resp)
        
        # Remove duplicates and sort by length (longer ones first, likely more detailed)
        responsibilities = list(set(responsibilities))
        responsibilities.sort(key=len, reverse=True)
        
        # Return top N most relevant responsibilities
        return responsibilities[:15] # Return a few more for better matching
        
    except Exception as e:
        logging.error(f"Failed to extract key responsibilities: {e}")
        return []

def parse_document(file_content: bytes, file_type: str, is_jd: bool) -> Dict:
    """Parse document and extract all relevant information"""
    try:
        # Extract text based on file type
        if file_type.lower() == 'pdf':
            raw_text = extract_text_from_pdf(file_content)
        elif file_type.lower() in ['docx', 'doc']:
            raw_text = extract_text_from_docx(file_content)
        elif file_type.lower() == 'txt':
            raw_text = extract_text_from_txt(file_content)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        if not raw_text:
            raise ValueError("No text could be extracted from the document")
        
        # Clean text for processing (for general text fields)
        cleaned_text = clean_text(raw_text)
        
        # Extract common information for both JD and CV
        parsed_data = {
            "text": raw_text,
            "cleaned_text": cleaned_text,
            "skills": extract_skills(raw_text), # Use raw text for skills to preserve casing for proper nouns in specific cases
            "experience": extract_experience(raw_text),
            "ctc": extract_ctc_info(raw_text)
        }
        
        # Extract CV-specific information
        if not is_jd:
            parsed_data["academic_info"] = extract_academic_info(raw_text)
            
        # Extract JD-specific information
        if is_jd:
            parsed_data["key_responsibilities"] = extract_key_responsibilities_from_jd(raw_text)
        
        return parsed_data
        
    except Exception as e:
        logging.error(f"Failed to parse document: {e}")
        # Re-raise the exception to be handled by the Flask app
        raise

def calculate_similarity_score(text1: str, text2: str) -> float:
    """Calculate similarity between two text snippets using spaCy's embeddings"""
    global nlp
    if not nlp:
        if not initialize_nlp():
            logging.warning("NLP model not loaded, falling back to basic word overlap for similarity.")
            # Fallback to simple word overlap if NLP is not available
            words1 = set(text1.lower().split())
            words2 = set(text2.lower().split())
            intersection = words1.intersection(words2)
            union = words1.union(words2)
            return len(intersection) / len(union) if union else 0

    try:
        doc1 = nlp(text1)
        doc2 = nlp(text2)
        # Ensure documents have vectors before calculating similarity
        if not doc1.has_vector or not doc2.has_vector:
             logging.warning("One or both documents do not have word vectors. Falling back to basic word overlap for similarity.")
             words1 = set(text1.lower().split())
             words2 = set(text2.lower().split())
             intersection = words1.intersection(words2)
             union = words1.union(words2)
             return len(intersection) / len(union) if union else 0

        return doc1.similarity(doc2)
    except Exception as e: # Catch more specific errors if known, otherwise general for robustness
        logging.error(f"Error calculating spaCy similarity, falling back to word overlap: {e}")
        # Fallback to simple word overlap
        words1 = set(text1.lower().split())
        words2 = set(text2.lower().split())
        intersection = words1.intersection(words2)
        union = words1.union(words2)
        return len(intersection) / len(union) if union else 0

def calculate_match_score(parsed_jd: Dict, parsed_cv: Dict) -> Dict:
    """Calculate comprehensive match score with detailed feedback"""
    try:
        result = {
            "score": 0,
            "matched_skills": [],
            "missing_jd_skills": [],
            "extra_cv_skills": [],
            "experience_feedback": "",
            "ctc_feedback": "",
            "academic_alignment_feedback": "",
            "jd_responsibilities_matched_in_cv": []
        }
        
        jd_skills = set(skill.lower() for skill in parsed_jd.get("skills", []))
        cv_skills = set(skill.lower() for skill in parsed_cv.get("skills", []))
        
        # Skills Analysis (60% weight) - Adjusted weight
        matched_skills = jd_skills.intersection(cv_skills)
        missing_skills = jd_skills - cv_skills
        extra_skills = cv_skills - jd_skills
        
        result["matched_skills"] = list(matched_skills)
        result["missing_jd_skills"] = list(missing_skills)
        result["extra_cv_skills"] = list(extra_skills)
        
        # Calculate skills score based on matched vs. required
        skills_score = 0
        if jd_skills:
            # Emphasize matching crucial skills more if available
            crucial_skills_jd = {s for s in jd_skills if s in PROFESSIONAL_SKILLS} # Skills from the defined list
            if crucial_skills_jd:
                matched_crucial = crucial_skills_jd.intersection(cv_skills)
                skills_score = (len(matched_crucial) / len(crucial_skills_jd)) * 100
            else: # If JD has no 'professional' skills, just use general skill overlap
                 skills_score = (len(matched_skills) / len(jd_skills)) * 100
        else:
            skills_score = 50 # Neutral if no skills defined in JD
        
        # Experience Analysis (20% weight)
        jd_exp = parsed_jd.get("experience", {})
        cv_exp = parsed_cv.get("experience", {})
        
        jd_years = jd_exp.get("years_of_experience", 0)
        cv_years = cv_exp.get("years_of_experience", 0)
        jd_seniority = jd_exp.get("seniority_level", "Entry-Level")
        cv_seniority = cv_exp.get("seniority_level", "Entry-Level")
        
        experience_score = 0
        # If JD specifies experience, compare
        if jd_years > 0:
            if cv_years >= jd_years:
                experience_score = 100
            elif cv_years >= jd_years * 0.9:
                experience_score = 90
            elif cv_years >= jd_years * 0.7:
                experience_score = 70
            elif cv_years >= jd_years * 0.5:
                experience_score = 50
            else:
                experience_score = max(0, (cv_years / jd_years) * 40) # Lower score for significant gaps
        else: # If JD doesn't specify years, a reasonable amount of experience is good
            if cv_years > 0:
                experience_score = 80 # Good if CV has experience but JD doesn't specify
            else:
                experience_score = 50 # Neutral if neither specifies
        
        # Adjust experience score based on seniority level alignment
        jd_seniority_val = SENIORITY_LEVELS.get(jd_seniority.lower(), 1)
        cv_seniority_val = SENIORITY_LEVELS.get(cv_seniority.lower(), 1)
        
        if cv_seniority_val >= jd_seniority_val:
            experience_score = min(100, experience_score + 10) # Bonus for matching or exceeding seniority
        elif cv_seniority_val < jd_seniority_val:
            experience_score = max(0, experience_score - 15) # Penalty for lower seniority

        # Experience feedback
        if cv_years >= jd_years and cv_seniority_val >= jd_seniority_val:
            result["experience_feedback"] = f"Excellent match: Candidate's {cv_years} years of {cv_seniority} experience matches or exceeds the JD's {jd_years}+ years {jd_seniority} requirement."
        elif cv_years >= jd_years * 0.7:
            result["experience_feedback"] = f"Good match: Candidate's {cv_years} years of {cv_seniority} experience is closely aligned with the JD's {jd_years}+ years {jd_seniority} requirement."
        else:
            result["experience_feedback"] = f"Experience gap: Candidate's {cv_years} years of {cv_seniority} experience is below the JD's {jd_years}+ years {jd_seniority} requirement. Consider for junior roles or if other areas compensate."
        
        # CTC Analysis (10% weight) - Adjusted weight
        jd_ctc = parsed_jd.get("ctc", {})
        cv_ctc = parsed_cv.get("ctc", {})
        
        ctc_score = 50 # Default neutral score if no info or major mismatch
        
        if jd_ctc and cv_ctc:
            jd_min = jd_ctc.get("min_value", 0)
            jd_max = jd_ctc.get("max_value", jd_min)
            cv_min = cv_ctc.get("min_value", 0)
            cv_max = cv_ctc.get("max_value", cv_min)
            
            # Simple check if currencies match
            if jd_ctc.get("currency") != cv_ctc.get("currency"):
                result["ctc_feedback"] = f"Warning: CTC currencies differ (JD: {jd_ctc.get('currency')}, CV: {cv_ctc.get('currency')}). Cannot compare directly."
                ctc_score = 20 # Low score if currencies don't match, as direct comparison is invalid
            else:
                currency_symbol = jd_ctc.get('currency', '') # Use the common currency symbol for feedback
                
                # Perfect overlap (CV expectation is within JD range)
                if cv_min >= jd_min and cv_max <= jd_max:
                    ctc_score = 100
                    result["ctc_feedback"] = f"Excellent alignment: Candidate's expected CTC ({currency_symbol} {cv_min:,.0f} - {cv_max:,.0f}) is perfectly within the JD's range ({currency_symbol} {jd_min:,.0f} - {jd_max:,.0f})."
                # CV expects less than JD min
                elif cv_max < jd_min:
                    ctc_score = 90
                    result["ctc_feedback"] = f"Favorable: Candidate's expected CTC ({currency_symbol} {cv_min:,.0f} - {cv_max:,.0f}) is below the JD's minimum ({currency_symbol} {jd_min:,.0f})."
                # Partial overlap (CV min is less than JD max, but CV max is higher than JD max)
                elif cv_min < jd_max and cv_max > jd_max:
                    # Calculate how far above the JD max the CV goes
                    # Using a simplified overlap percentage for scoring clarity
                    overlap_amount = min(cv_max, jd_max) - max(cv_min, jd_min)
                    if overlap_amount > 0:
                        overlap_ratio = overlap_amount / (jd_max - jd_min) if (jd_max - jd_min) > 0 else 0
                        ctc_score = 50 + (overlap_ratio * 50) # Scale score based on overlap
                    else:
                        ctc_score = 50 # No positive overlap
                    
                    result["ctc_feedback"] = f"Moderate overlap: Candidate's expected CTC ({currency_symbol} {cv_min:,.0f} - {cv_max:,.0f}) partially overlaps with JD range ({currency_symbol} {jd_min:,.0f} - {jd_max:,.0f}), leaning slightly higher."
                # CV expects more than JD max
                elif cv_min > jd_max:
                    gap_percentage = ((cv_min - jd_max) / jd_max) * 100 if jd_max > 0 else 100
                    if gap_percentage <= 15:
                        ctc_score = 70
                        result["ctc_feedback"] = f"Slight mismatch: Candidate's expected CTC ({currency_symbol} {cv_min:,.0f} - {cv_max:,.0f}) is slightly above the JD's maximum ({currency_symbol} {jd_max:,.0f}). May be negotiable."
                    else:
                        ctc_score = 30
                        result["ctc_feedback"] = f"Significant mismatch: Candidate's expected CTC ({currency_symbol} {cv_min:,.0f} - {cv_max:,.0f}) is significantly above the JD's maximum ({currency_symbol} {jd_max:,.0f}). Low alignment."
                else: # Any other unexpected case, default to neutral
                    result["ctc_feedback"] = "CTC comparison inconclusive due to complex ranges."

        else: # One or both CTCs are missing
            result["ctc_feedback"] = "CTC information not available for one or both documents for comparison."
            ctc_score = 50 # Neutral score if info is missing

        # JD Responsibilities Matching (5% weight) - Adjusted weight
        jd_responsibilities = parsed_jd.get("key_responsibilities", [])
        cv_text_raw = parsed_cv.get("text", "") # Use raw text for responsibility matching
        
        responsibilities_match_score = 0
        matched_responsibilities_details = []
        
        if jd_responsibilities and cv_text_raw:
            individual_scores = []
            for responsibility in jd_responsibilities:
                found_in_cv = False
                relevant_snippet = ""
                highest_snippet_score = 0.0 # Initialize as float
                
                # Check for direct phrase match first
                if responsibility.lower() in cv_text_raw.lower():
                    found_in_cv = True
                    # Find and extract a relevant sentence or snippet
                    for sentence in sent_tokenize(cv_text_raw):
                        if responsibility.lower() in sentence.lower():
                            relevant_snippet = sentence[:200] + "..." if len(sentence) > 200 else sentence
                            break # Found a direct match sentence
                    highest_snippet_score = 1.0 # Max confidence for direct match
                    individual_scores.append(100)
                else:
                    # Use semantic similarity (if NLP is loaded and has vectors) or keyword overlap as fallback
                    if nlp and nlp.vocab.vectors.name: # Check if NLP has vectors loaded
                        jd_resp_doc = nlp(responsibility)
                        cv_sentences = sent_tokenize(cv_text_raw)
                        
                        for sentence in cv_sentences:
                            cv_sent_doc = nlp(sentence)
                            if jd_resp_doc.has_vector and cv_sent_doc.has_vector:
                                similarity = jd_resp_doc.similarity(cv_sent_doc)
                                if similarity > highest_snippet_score:
                                    highest_snippet_score = similarity
                                    relevant_snippet = sentence[:200] + "..." if len(sentence) > 200 else sentence
                                    
                        if highest_snippet_score >= 0.7: # Threshold for strong semantic match
                            found_in_cv = True
                            individual_scores.append(highest_snippet_score * 100)
                        elif highest_snippet_score > 0.3: # Partial semantic match
                             individual_scores.append(highest_snippet_score * 50) # Give partial credit
                        else:
                            individual_scores.append(0) # No significant semantic match
                    else: # Fallback to keyword overlap if NLP not available or no vectors
                        responsibility_words = set(clean_text(responsibility).split())
                        cv_words_cleaned = set(clean_text(cv_text_raw).split())
                        
                        common_words_count = len(responsibility_words.intersection(cv_words_cleaned))
                        if responsibility_words:
                            keyword_match_ratio = common_words_count / len(responsibility_words)
                        else:
                            keyword_match_ratio = 0
                        
                        if keyword_match_ratio >= 0.4: # Adjustable threshold
                            found_in_cv = True
                            individual_scores.append(keyword_match_ratio * 100)
                            # Find relevant sentence via keyword overlap
                            cv_sentences = sent_tokenize(parsed_cv.get("text", ""))
                            best_sentence = ""
                            best_overlap_score = 0
                            for sentence in cv_sentences:
                                sentence_words = set(clean_text(sentence).split())
                                overlap = len(sentence_words.intersection(responsibility_words))
                                if overlap > best_overlap_score:
                                    best_overlap_score = overlap
                                    best_sentence = sentence
                            relevant_snippet = best_sentence[:200] + "..." if len(best_sentence) > 200 else best_sentence
                        else:
                            individual_scores.append(0)
                
                matched_responsibilities_details.append({
                    "responsibility": responsibility,
                    "found_in_cv": found_in_cv,
                    "relevant_snippet": relevant_snippet,
                    "confidence_score": round(highest_snippet_score * 100, 2) if nlp else round(individual_scores[-1], 2) # Use last appended score
                })
            
            if individual_scores:
                responsibilities_match_score = sum(individual_scores) / len(individual_scores)
            else:
                responsibilities_match_score = 0 # No responsibilities to compare
        else:
            responsibilities_match_score = 50 # Neutral if JD has no responsibilities
        
        result["jd_responsibilities_matched_in_cv"] = matched_responsibilities_details
        
        # Academic Alignment (5% weight)
        academic_score = 50 # Neutral if no academic info or not a major factor
        cv_academic = parsed_cv.get("academic_info", {})
        academic_feedback_list = [] # Renamed to avoid conflict with result key
        
        if cv_academic.get("degrees"):
            academic_feedback_list.append(f"Education: {len(cv_academic['degrees'])} degree(s) found.")
            academic_score += 10 # Bonus for degrees
        
        if cv_academic.get("universities"):
             academic_feedback_list.append(f"University mentions: {len(cv_academic['universities'])}.")
             # Further logic could compare university reputation etc.
        
        if cv_academic.get("publications"):
            academic_feedback_list.append(f"Research: {len(cv_academic['publications'])} publication(s).")
            academic_score += 15 # Higher bonus for publications
        
        if cv_academic.get("awards"):
            academic_feedback_list.append(f"Recognition: {len(cv_academic['awards'])} award(s)/achievement(s).")
            academic_score += 10 # Bonus for awards
        
        result["academic_alignment_feedback"] = " ".join(academic_feedback_list).strip() or "Limited academic information available."
        academic_score = min(100, academic_score) # Cap at 100

        # Calculate final weighted score
        final_score = (skills_score * 0.60) + \
                      (experience_score * 0.20) + \
                      (ctc_score * 0.10) + \
                      (responsibilities_match_score * 0.05) + \
                      (academic_score * 0.05)
        
        result["score"] = round(final_score, 1)
        
        return result
        
    except Exception as e:
        logging.error(f"Failed to calculate match score: {e}")
        # Return a structured error response in case of a critical failure during scoring
        return {
            "score": 0,
            "matched_skills": [],
            "missing_jd_skills": [],
            "extra_cv_skills": [],
            "experience_feedback": f"Error calculating experience match: {e}",
            "ctc_feedback": f"Error calculating CTC match: {e}",
            "academic_alignment_feedback": f"Error calculating academic alignment: {e}",
            "jd_responsibilities_matched_in_cv": [],
            "overall_error": "Failed to calculate comprehensive match score."
        }

# Initialize NLP on module import.
# This ensures spaCy and NLTK are ready when ats_core.py is first imported.
initialize_nlp()